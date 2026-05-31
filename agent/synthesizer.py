import os
from openai import AsyncOpenAI
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from config import OPENROUTER_API_KEY, BASE_URL, MODEL

client = AsyncOpenAI(
    api_key=OPENROUTER_API_KEY,
    base_url=BASE_URL
)

def save_as_docx(query: str, report_text: str, filepath: str):
    """Convert report text into a formatted Word document."""

    doc = Document()

    # Title
    title = doc.add_heading("Deep Research Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Query
    doc.add_heading("Research Query", level=1)
    doc.add_paragraph(query)

    # Divider
    doc.add_paragraph("─" * 60)

    # Report content — split by lines
    for line in report_text.split("\n"):
        line = line.strip()
        if not line:
            continue

        # Detect markdown headings → Word headings
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)

        # Detect bullet points
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")

        # Detect numbered list
        elif line[0].isdigit() and line[1:3] in (". ", ") "):
            doc.add_paragraph(line[3:], style="List Number")

        # Normal paragraph
        else:
            doc.add_paragraph(line)

    doc.save(filepath)


async def synthesize_report(query: str, findings_dir: str = "findings") -> str:
    """Read all findings and compile into a final research report."""

    all_findings = []

    if os.path.exists(findings_dir):
        for filename in os.listdir(findings_dir):
            filepath = os.path.join(findings_dir, filename)
            with open(filepath, "r", encoding="utf-8") as f:
                content = f.read()
                all_findings.append(f"--- {filename} ---\n{content}")

    if not all_findings:
        return "No findings to synthesize."

    combined = "\n\n".join(all_findings)

    # Final LLM call
    response = await client.chat.completions.create(
        model=MODEL,
        max_tokens=2000,
        messages=[
            {
                "role": "system",
                "content": (
                    "You are a research analyst. "
                    "Given raw research findings, write a clean, structured, "
                    "insightful report with headings and bullet points. "
                    "Use ## for main sections, ### for subsections, - for bullets. "
                    "Write ONLY from the provided findings. "
                    "If information is not in findings, say 'not found in sources'. "
                    "Never invent statistics, company names, or quotes."
                )
            },
            {
                "role": "user",
                "content": f"Query: {query}\n\nFindings:\n{combined}"
            }
        ]
    )

    return response.choices[0].message.content